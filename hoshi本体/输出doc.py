import logging

import cv2
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import 表格识别
from . import 文字提取


def 排序键(x):
    if isinstance(x, 表格识别.表格):
        return x.原位置['top']
    return x['top']


def 输出页(document, 目录信息, 段落信息, 表格组, 图块组):
    for i, x in enumerate(sorted(段落信息 + 表格组 + 图块组 + 目录信息, key=lambda i: 排序键(i))):
        
        if x in 表格组:
            table = document.add_table(rows=x.尺寸[0], cols=x.尺寸[1], style='Table Grid')
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i in range(x.尺寸[0]):
                for j in range(x.尺寸[1]):
                    if x.格内容[i][j]:
                        table.cell(i, j).text = x.格内容[i][j]
                    w = x.格范围(i, j)
                    if w and w != (i, j):
                        table.cell(i, j).merge(table.cell(*w))
            document.add_paragraph('')
        
        elif x in 目录信息:
            段落 = x
            document.add_paragraph(f'(目录){x["内容"]}')
            
        elif x in 段落信息:
            段落 = x
            if 段落['样式'] == '居中':
                总字 = '\n'.join([
                    i['内容']
                    for i in 段落['行组']
                ])
                document.add_paragraph(总字)
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                总字 = '\n'.join([
                    int(i['缩进'] / ((i['bottom']-i['top']) * 2)) * ' ' + i['内容']
                    for i in 段落['行组']
                ])
                document.add_paragraph(总字)
                
        elif x in 图块组:
            表 = x
            cv2.imwrite(f'./_temp/_.jpg', 表['内容'])
            document.add_picture(f'./_temp/_.jpg', width=Inches((表['right'] - 表['left']) / 600))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            0/0

    document.add_page_break()


def 输出(文件名, 页组):
    document = Document()
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal'].font.size = Pt(8)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for 页 in 页组:
        输出页(document, 页['目录信息'], 页['段落信息'], 页['表格组'], 页['图块组'])

    document.save(文件名)

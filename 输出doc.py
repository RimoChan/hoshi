import logging

import cv2
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def 输出页(document, 段落信息, 表格组, 图块组):
    for i, x in enumerate(sorted(段落信息 + 表格组 + 图块组, key=lambda i: i['top'])):
        if '行组' in x:
            段落 = x
            if 段落['样式'] == '居中':
                总字 = '\n'.join([
                    ''.join([i for i, _ in i['内容']])
                    for i in 段落['行组']
                ])
                document.add_paragraph(总字)
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                总字 = '\n'.join([
                    int(i['缩进'] / (i['height'] * 2)) * ' ' + ''.join([i for i, _ in i['内容']])
                    for i in 段落['行组']
                ])
                document.add_paragraph(总字)
        else:
            表 = x
            cv2.imwrite(f'./_temp/_.jpg', 表['内容'])
            document.add_picture(f'./_temp/_.jpg', width=Inches((表['right'] - 表['left']) / 600))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()


def 输出(文件名, 页组):
    document = Document()
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal'].font.size = Pt(8)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for 页 in 页组:
        输出页(document, 页['段落信息'], 页['表格组'], 页['图块组'])

    document.save(文件名)
